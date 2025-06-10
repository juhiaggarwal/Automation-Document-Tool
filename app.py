import re
import random
import smtplib
import os
from email.mime.text import MIMEText
from docx.shared import RGBColor
import requests
import uuid
from docx.shared import Pt
from flask import Flask, render_template, request, jsonify, redirect, url_for, session
from flask import Flask, request, jsonify
from flask_mail import Mail, Message
from flask_sqlalchemy import SQLAlchemy
from flask import render_template, url_for
from flask_login import login_required, current_user
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pytz import timezone
from datetime import datetime
from flask_cors import CORS
from dotenv import load_dotenv
from flask import send_file, request
from docx2pdf import convert
from werkzeug.security import generate_password_hash, check_password_hash
from flask_dance.contrib.google import make_google_blueprint, google
from flask_dance.contrib.github import make_github_blueprint, github
from flask_login import LoginManager, login_required, login_user, logout_user, current_user
from flask import Flask, request, send_from_directory, jsonify, render_template, Response
from markupsafe import escape
import tempfile
from docx.shared import Inches
from flask_cors import CORS
from flask import render_template, request, redirect, url_for, session
from flask_login import login_required, current_user
from models import db, User, Submission  # ensure these are imported properly
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader
from PyPDF2 import PdfMerger
from docx import Document as DocxDocument
from pdf2image import convert_from_path
from docx.shared import Inches,RGBColor
import io
from pdf2docx import Converter
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import zipfile
from datetime import datetime
from werkzeug.utils import secure_filename
import pdfkit
from config import Config
from models import db,User,UploadedFile
from flask import Flask, request, send_file, after_this_request
import json
from flask import request, jsonify, send_file
from flask_login import login_required, current_user
from models import Submission, User 
import subprocess
from asposewordscloud import Configuration, WordsApi
from asposewordscloud.models import PdfSaveOptionsData
import fitz  # PyMuPDF
from flask import request, jsonify, flash, redirect, url_for, render_template
from flask_login import login_required, logout_user, current_user
from models import db, User  # Ensure your User model is imported
from docx.shared import Inches
import tempfile




# Allow OAuth in insecure (non-HTTPS) for testing
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

# Load environment variables
load_dotenv()   

# Initialize Flask app and configs
app = Flask(__name__)
app.config.from_object(Config)

# Assume you loaded these into Config:
# EMAIL_USER, EMAIL_PASS, EMAIL_SERVER, SMTP_PORT
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False
app.config['MAIL_USERNAME'] = Config.EMAIL_USER
app.config['MAIL_PASSWORD'] = Config.EMAIL_PASS
app.config['MAIL_DEFAULT_SENDER'] = Config.EMAIL_USER

mail = Mail(app)

#SAVE FOLDER LOCATION
UPLOAD_DIR = tempfile.gettempdir()
#UPLOAD_DIR = os.path.join(os.getcwd(), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


db.init_app(app)
CORS(app)

#cloudinary
def upload_to_cloudinary(local_file_path):
    try:
        result = cloudinary.uploader.upload(local_file_path)
        return result["secure_url"]
    except Exception as e:
        print("Cloudinary upload failed:", e)
        return None





#FAQ
# Predefined app-specific FAQ responses with emojis
app_faq = {
    "how to register": "📝 You can register in three ways:\n1️⃣ Continue with Google – Sign in using your Gmail account.\n2️⃣ OTP verification via email – Enter your email, receive an OTP, and verify it to register.\n 3.Continue with Github –Sign in using our GitHub Account.",
    "password error": "🔐 Make sure your password includes:\n- One uppercase letter\n- One lowercase letter\n- One number\n- One special character (@#$%^&+=)\n- At least 8 characters",
    "how to login": "🔓 Go to Login Page after successful registration, enter your email and password, then hit Login.",
    "how to generate document": "📊 After successful registration and login, go to the Dashboard and click on Create a Documentary.",
    "how to logout": "🚪 Click the Logout button or open your Profile Icon and select Logout.",
    "how to edit profile": "⚙️ In the Dashboard, click the Profile Icon, then choose Edit Profile.",
    "how to preview document": "👀 Navigate to the Automation Documentation page and click Preview to review your documentary.",
    "how to submit document": "✅ After generating your documentary on the Automation Documentation Page, scroll to the bottom and click Submit.",
    "how to export to pdf": "📄 First generate your document fill the data you want to add in the documetary and hit the submit button  , then click the Export to PDF button to download it as a PDF file."
}


# Setup Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# ------------------ OAuth Setup ------------------

# Google OAuth
if app.config.get("GOOGLE_CLIENT_ID") and app.config.get("GOOGLE_CLIENT_SECRET"):
    google_bp = make_google_blueprint(
        client_id=app.config["GOOGLE_CLIENT_ID"],
        client_secret=app.config["GOOGLE_CLIENT_SECRET"],
        scope=["openid", "https://www.googleapis.com/auth/userinfo.email", 
               "https://www.googleapis.com/auth/userinfo.profile"],
        redirect_to="google_auth_complete"
    )
    app.register_blueprint(google_bp, url_prefix="/login")
else:
    print("⚠️ Missing Google OAuth configuration in .env")

# GitHub OAuth
if app.config.get("GITHUB_CLIENT_ID") and app.config.get("GITHUB_CLIENT_SECRET"):
    github_bp = make_github_blueprint(
        client_id=app.config["GITHUB_CLIENT_ID"],
        client_secret=app.config["GITHUB_CLIENT_SECRET"],
        redirect_to="github_auth_complete"
    )
    app.register_blueprint(github_bp, url_prefix="/login")
else:
    print("⚠️ Missing GitHub OAuth configuration in .env")

# Create tables
with app.app_context():
    db.create_all()

# ------------------ Routes ------------------

@app.route('/')
def index():
    return redirect(url_for('login'))

@app.route('/register', methods=['GET'])
def register():
    return render_template('register.html')

@app.route('/login', methods=['GET'])
def login():
 
    return render_template('login.html')

#-----------------DASHBOARD--------------------

@app.route('/dashboard')
@login_required
def dashboard():
    user = User.query.get(session.get('user_id'))
    if not user:
        return redirect(url_for('login'))

    first_name = user.first_name or user.email.split('@')[0].capitalize()
    return render_template('dashboard.html', first_name=first_name)
 
# ------------------ OTP ------------------

otp_storage = {}

@app.route('/send-otp', methods=['POST'])
def send_otp():
    data = request.get_json()
    email = data.get('email')
    if not email:
        return jsonify({'message': 'Email is required'}), 400

    otp = str(random.randint(100000, 999999))
    otp_storage[email] = otp

    msg = MIMEText(f'Your OTP is {otp}')
    msg['Subject'] = 'OTP Verification'
    msg['From'] = app.config['EMAIL_USER']
    msg['To'] = email

    try:
        with smtplib.SMTP(app.config['SMTP_SERVER'], app.config['SMTP_PORT']) as server:
            server.starttls()
            server.login(app.config['EMAIL_USER'], app.config['EMAIL_PASS'])
            server.send_message(msg)
        return jsonify({'message': f'OTP sent to {email}'})
    except Exception as e:
        return jsonify({'message': 'Failed to send OTP', 'error': str(e)}), 500

@app.route('/verify-otp', methods=['POST'])
def verify_otp():
    data = request.get_json()
    email = data.get('email')
    otp = data.get('otp')
    if otp_storage.get(email) == otp:
        return jsonify({'message': 'OTP verified'}), 200
    return jsonify({'message': 'Invalid OTP'}), 400

#-----dashboard---------------
@app.route('/create-report', methods=['GET'])
@login_required
def create_report():
    return render_template('1.html')



#----------LOGOUT--------------
@app.route('/logout', methods=['GET', 'POST'])
@login_required
def logout():
    if request.method == "POST":
        if request.is_json:
            data = request.get_json()
            feedback = data.get("feedback")
            rating = data.get("rating")
            do_logout = data.get("logout")  # NEW: distinguish logout trigger

            if do_logout:
                logout_user()
                return jsonify({"message": "Logged out successfully"}), 200

            if feedback and rating:
                print(f"User: {current_user.email}, Feedback: {feedback}, Rating: {rating}")

                if current_user.feedback_given is None:
                    current_user.feedback_given = 0
                current_user.feedback_given += 1

                db.session.commit()
                return jsonify({"message": "Feedback saved"}), 200

            return jsonify({"error": "Invalid request"}), 400

        # Optional: legacy form fallback
        feedback = request.form.get("feedback")
        rating = request.form.get("rating")
        if feedback and rating:
            if current_user.feedback_given is None:
                current_user.feedback_given = 0
            current_user.feedback_given += 1
            db.session.commit()

        logout_user()
        flash("You have been logged out. Thanks for your feedback!", "info")
        return redirect(url_for("login"))

    return render_template("logout.html")





#---------------PROFILE-------------------------


@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    # Get current user object
    user = User.query.get(current_user.id)

    if not user:
        return redirect(url_for('login'))

    # Handle POST: update name
    if request.method == 'POST':
        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')

        # Save changes
        user.first_name = first_name
        user.last_name = last_name
        db.session.commit()

        return redirect(url_for('profile'))

    # For GET: calculate user's report and feedback counts
    reports_count = len(user.submissions)
    feedbacks_count = user.feedback_given or 0  # or track via separate table if needed

    return render_template(
        'profile.html',
        first_name=user.first_name or '',
        last_name=user.last_name or '',
        email=user.email,
        reports_count=reports_count,
        feedbacks_count=feedbacks_count
    )





# ------------------ Registration ------------------

def is_strong_password(password):
    return (
        len(password) >= 8 and
        re.search(r'[A-Z]', password) and
        re.search(r'[a-z]', password) and
        re.search(r'[0-9]', password) and
        re.search(r'[@#$%^&+=]', password)
    )

@app.route('/register', methods=['POST'])
def register_user():
    form = request.form
    email = form.get('email')
    otp = form.get('otp')
    password = form.get('password')
    first_name = form.get('first_name')
    last_name = form.get('last_name')

    # ✅ OTP validation
    if otp_storage.get(email) != otp:
        return jsonify({'message': 'Invalid OTP'}), 400

    # ✅ Password strength check
    if not is_strong_password(password):
        return jsonify({
            'message': 'Password not strong. Make sure your password includes:\n'
                       '- One uppercase letter\n'
                       '- One lowercase letter\n'
                       '- One number\n'
                       '- One special character (@#$%^&+=)\n'
                       '- At least 8 characters'
        }), 400

    # ✅ Check for duplicate email
    if User.query.filter_by(email=email).first():
        return jsonify({'message': 'User already registered! You can Login!!'}), 400

    # ✅ Store hashed + plain password
    hashed_password = generate_password_hash(password)
    user = User(
        first_name=first_name,
        last_name=last_name,
        email=email,
        password=hashed_password,
        original_password=password  # ❌ FIXED: was ['password'] (a list), now correct
    )
    db.session.add(user)
    db.session.commit()

    # ✅ Send confirmation email
    msg = MIMEText(f"Thank you for registering!\n\n📧 Email: {email}\n🔐 Password: {password}")
    msg['Subject'] = 'Registration Successful'
    msg['From'] = app.config['EMAIL_USER']
    msg['To'] = email

    try:
        with smtplib.SMTP(app.config['SMTP_SERVER'], app.config['SMTP_PORT']) as server:
            server.starttls()
            server.login(app.config['EMAIL_USER'], app.config['EMAIL_PASS'])
            server.send_message(msg)
    except Exception as e:
        return jsonify({'message': 'Registered, but email failed', 'error': str(e)}), 500

    return jsonify({'message': 'Registration successful'}), 200


# ------------------ Login ------------------

@app.route('/login', methods=['POST'])
def login_user_route():
    form = request.form
    email = form.get('email')
    password = form.get('password')

    user = User.query.filter_by(email=email).first()
    if not user or not check_password_hash(user.password, password):
        return jsonify({'message': 'Invalid credentials'}), 401

    session['user_id'] = user.id
    session['user_email'] = user.email
    login_user(user)
    return jsonify({'message': 'Login successful'}), 200

    
    
# ---------------- Password Reset -----------------------
@app.route("/reset-password", methods=["POST"])
def reset_password():
    data = request.json
    email = data.get("email")

    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify({"success": False, "message": "This email is not registered."}), 400

    full_name = f"{user.first_name or ''} {user.last_name or ''}".strip()
    subject = "Password Reset - Your Registered Credentials"
    body = f"""Hello {full_name or 'User'},

You have requested a password reset.

Here are your registration details:
📧 Email: {user.email}
🔐 Password: {user.original_password}

If you did not request this, please ignore this email.

Regards,  
Support Team
"""

    try:
        msg = Message(subject, sender=app.config['MAIL_USERNAME'], recipients=[email])
        msg.body = body
        mail.send(msg)
        return jsonify({"success": True, "message": "Reset email sent successfully."})
    except Exception as e:
        return jsonify({"success": False, "message": f"Failed to send email: {str(e)}"}), 500

    


#------------REPORTED GENERATED----------------
from flask import render_template
from flask_login import login_required, current_user
from pytz import timezone, utc

@app.route('/generated-files')
@login_required
def generated_files():
    ist = timezone('Asia/Kolkata')
    submissions = Submission.query.filter_by(user_id=current_user.id).order_by(Submission.timestamp.desc()).all()

    # Convert naive timestamps (no timezone) to UTC, then to IST
    for s in submissions:
        if s.timestamp.tzinfo is None:
            s.timestamp = utc.localize(s.timestamp)  # Treat as UTC if naive
        s.timestamp = s.timestamp.astimezone(ist)     # Convert to IST

    return render_template('generatedreports.html',
                           submissions=submissions,
                           username=current_user.first_name or current_user.email.split('@')[0])






 



# ------------------ Google Auth ------------------

@app.route("/google_auth_complete")
def google_auth_complete():
    if not google.authorized:
        return redirect(url_for("google.login"))
    resp = google.get("/oauth2/v2/userinfo")
    if not resp.ok:
        return jsonify({"error": "Failed to fetch user info"}), 500

    user_data = resp.json()
    email = user_data["email"]
    user = User.query.filter_by(email=email).first()
    if not user:
        user = User(first_name=user_data.get("given_name", ""), last_name=user_data.get("family_name", ""),
                    email=email, password=generate_password_hash("google"))
        db.session.add(user)
        db.session.commit()
    login_user(user)
    session['user_email'] = email
    session['user_id'] = user.id
    return redirect(url_for("dashboard"))

# ------------------ GitHub Auth ------------------

@app.route("/github_auth_complete")
def github_auth_complete():
    if not github.authorized:
        return redirect(url_for("github.login"))
    resp = github.get("/user")
    if not resp.ok:
        return jsonify({"error": "Failed to fetch GitHub user info"}), 500

    user_data = resp.json()
    email = user_data.get("email") or f'{user_data["login"]}@github.com'
    user = User.query.filter_by(email=email).first()
    if not user:
        user = User(first_name=user_data["login"], last_name="", email=email,
                    password=generate_password_hash("github"))
        db.session.add(user)
        db.session.commit()
    login_user(user)
    session['user_email'] = email
    session['user_id'] = user.id
    return redirect(url_for("dashboard"))

# ------------------ Chat Assistant ------------------
@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    query = data.get("query", "")
    if not query:
        return jsonify({'error': 'No query provided'}), 400

    normalized_query = query.lower().strip()

    # Check predefined FAQ answers
    for keyword, answer in app_faq.items():
        if keyword in normalized_query:
            return jsonify({"response": answer})

    # Fallback to Gemini API
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
    if not GEMINI_API_KEY:
        return jsonify({'error': 'Missing Gemini API key'}), 500

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}"
    headers = {"Content-Type": "application/json"}
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {
                        "text": query
                    }
                ]
            }
        ],
        "generationConfig": {
            "temperature": 0.7,
            "topK": 40,
            "topP": 0.95,
            "maxOutputTokens": 1024
        }
    }

    try:
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()
        result = response.json()

        reply = result["candidates"][0]["content"]["parts"][0]["text"]
        return jsonify({"response": reply})

    except requests.exceptions.RequestException as e:
        print("Gemini API request failed:", e.response.text if e.response else str(e))
        return jsonify({'error': 'Gemini API request failed'}), 500

    except Exception as e:
        print("Internal Server Error:", str(e))
        return jsonify({'error': 'Internal server error'}), 500



# ------------------ Create Report ------------------

#----------------DOCX-------------------


def embed_pdf_as_images(doc, pdf_path):
    try:
        pages = convert_from_path(pdf_path)
        for i, page in enumerate(pages):
            doc.add_paragraph(f"Page {i+1}")
            img_path = f"{pdf_path}_page_{i}.png"
            page.save(img_path, 'PNG')
            doc.add_picture(img_path, width=Inches(5.5))
            os.remove(img_path)  # Clean up temp image
    except Exception as e:
        doc.add_paragraph(f"(Error embedding PDF as images: {e})")


def embed_docx(target_doc, filepath):
    try:
        source_doc = DocxDocument(filepath)
        for para in source_doc.paragraphs:
            target_doc.add_paragraph(para.text)
    except Exception as e:
        target_doc.add_paragraph(f"(Error embedding DOCX: {e})")
        
def embed_text_file(target_doc, filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                target_doc.add_paragraph(line.strip())
    except Exception as e:
        target_doc.add_paragraph(f"(Error reading TXT file: {e})")


def pdf_to_docx_in_memory(pdf_path):
    """Convert PDF file to DOCX document object in memory."""
    temp_docx_path = pdf_path + ".converted.docx"
    cv = Converter(pdf_path)
    cv.convert(temp_docx_path, start=0, end=None)
    cv.close()
    doc = Document(temp_docx_path)
    os.remove(temp_docx_path)  # clean up temp file
    return doc


def embed_txt_into_doc(target_doc, txt_path):
    """Embed text file content into doc."""
    with open(txt_path, 'r', encoding='utf-8') as f:
        for line in f:
            target_doc.add_paragraph(line.rstrip('\n'))


#------------- PREVIEW----------------

@app.route('/preview', methods=['POST'])
def preview():
    import os
    html_preview = """
    <style>
    table {width:100%;border-collapse:collapse;}
    th,td {border:1px solid #333;padding:8px;text-align:left;}
    th {background-color:#f5f5f5;}
    </style>
    <h2>Preview Responses</h2>
    """

    # Use the correct upload directory
    UPLOAD_DIR = tempfile.gettempdir()
    os.makedirs(UPLOAD_DIR, exist_ok=True)

    for key in request.form:
        if key.startswith("question_"):
            q_no = key.split("_")[1]
            answer = request.form.get(key)
            label = escape(request.form.get(f"label_{q_no}", f"Question {q_no}"))

            html_preview += f"<div class='question-block'><div class='question-title'>Q{q_no}: {label}</div>"

            if answer == "skip":
                html_preview += "<div class='skipped'>[Skipped]</div>"
            elif answer.startswith("table:"):
                try:
                    rows, cols = map(int, answer.split(":")[1].split("x"))
                    html_preview += "<table><thead><tr>" + "".join(f"<th>Col {c+1}</th>" for c in range(cols)) + "<th>Uploads</th></tr></thead><tbody>"
                    for r in range(rows):
                        html_preview += "<tr>" + "".join(f"<td>Row {r+1} Col {c+1}</td>" for c in range(cols))
                        file_found = False
                        for key_prefix in [f"file_q{q_no}_r{r}", f"q{q_no}_r{r}_upload"]:
                            file = request.files.get(key_prefix)
                            if file and file.filename:
                                filename = secure_filename(file.filename)
                                filepath = os.path.join(UPLOAD_DIR, filename)
                                file.save(filepath)
                                file_url = f"/uploads/{filename}"
                                if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                                    html_preview += f"<td><img src='{file_url}' style='max-height:100px;'></td>"
                                else:
                                    html_preview += f"<td><a href='{file_url}' target='_blank'>{filename}</a></td>"
                                file_found = True
                                break
                        if not file_found:
                            html_preview += "<td>No file</td>"
                        html_preview += "</tr>"
                    html_preview += "</tbody></table>"
                except Exception as e:
                    html_preview += f"<div class='skipped'>Invalid table format: {e}</div>"
            else:
                html_preview += f"<div>{escape(answer)}</div>"

            # Single files (non-table)
            for file_key in [f"file_question_{q_no}", f"q{q_no}_upload"]:
                file = request.files.get(file_key)
                if file and file.filename:
                    filename = secure_filename(file.filename)
                    filepath = os.path.join(UPLOAD_DIR, filename)
                    file.save(filepath)
                    file_url = f"/uploads/{filename}"
                    if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                        html_preview += f"<div><img src='{file_url}' style='max-height:150px;'></div>"
                    else:
                        html_preview += f"<div><a href='{file_url}' target='_blank'>{filename}</a></div>"
                    break

            html_preview += "</div>"

    return Response(html_preview, mimetype="text/html")

                            
                            
#-------------SUBMIT----------------
# --------------------------------------
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '008000')  # green
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    new_run.append(rPr)

    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# --------------------------------------
# Extract text from PDF
# --------------------------------------
def pdf_to_docx_in_memory(pdf_path):
    doc = Document()
    try:
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
    except Exception as e:
        doc.add_paragraph(f"(Error reading PDF: {e})")
    return doc

# --------------------------------------
# Convert PDF pages to images and add to docx
# --------------------------------------

def pdf_to_images_with_fitz(pdf_path, doc):
    try:
        pdf = fitz.open(pdf_path)
        for i, page in enumerate(pdf):
            pix = page.get_pixmap(dpi=150)  # Good quality
            temp_path = os.path.join(UPLOAD_DIR, f"page_{i + 1}.png")
            pix.save(temp_path)

            doc.add_paragraph(f"Page {i + 1}")
            doc.add_picture(temp_path, width=Inches(5))

            os.remove(temp_path)
        pdf.close()
    except Exception as e:
        print(f"PDF to image failed: {e}")
        doc.add_paragraph(f"(Error converting PDF to images: {e})")



# --------------------------------------
# Submit route
# --------------------------------------
@app.route('/submit', methods=['POST'])
def submit():
    try:
        form = request.form
        files = request.files

        username = form.get('username', 'Unknown User')
        filename = form.get('filename', 'submission')

        response_path = os.path.join(UPLOAD_DIR, f"{filename}_response.docx")
        uploads_path = os.path.join(UPLOAD_DIR, f"{filename}_uploads.docx")

        response_doc = Document()
        response_doc.add_heading(f"Response of {username}", 0)

        upload_doc = Document()
        upload_doc.add_heading(f"Uploads of {username}", 0)

        # Extract question IDs in order
        qids = []
        seen = set()
        for key in form:
            if key.endswith("_question"):
                qid = key.split("_")[0]
                if qid not in seen:
                    seen.add(qid)
                    qids.append(qid)

        for qid in qids:
            question = form.get(f"{qid}_question", "Untitled Question")
            response_type = form.get(f"{qid}_type", "unknown")

            response_doc.add_heading(f"Question: {question}", level=1)
            upload_doc.add_heading(f"Question: {question}", level=1)

            if response_type == 'text':
                response = form.get(f"{qid}_response", "No response provided.")
                response_doc.add_paragraph(response)
                upload_doc.add_paragraph("No file uploaded because only text was provided.")

            elif response_type == 'table':
                headers = eval(form.get(f"{qid}_table_headers", "[]"))
                if not headers:
                    response_doc.add_paragraph("No valid table headers or data.")
                    upload_doc.add_paragraph("No file uploaded because no valid table found.")
                    continue

                table = response_doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h

                any_file_uploaded = False
                max_rows = 50  # Adjust as needed

                for row_index in range(max_rows):
                    row_has_file = any(f"{qid}_r{row_index}_c{i}_file" in files for i in range(len(headers)))
                    row_data = [form.get(f"{qid}_r{row_index}_c{i}", "") for i in range(len(headers))]

                    # Add row even if only files are uploaded
                    if not any(row_data) and not row_has_file:
                        continue


                    row = table.add_row().cells
                    for col_index, val in enumerate(row_data):
                        para = row[col_index].paragraphs[0]
                        if val:
                            para.add_run(val)

                        file_key = f"{qid}_r{row_index}_c{col_index}_file"
                        name_key = f"{qid}_r{row_index}_c{col_index}_filename"

                        if file_key in files:
                            any_file_uploaded = True
                            file = files[file_key]
                            original_name = secure_filename(form.get(name_key) or file.filename)
                            unique_id = uuid.uuid4().hex
                            fname = f"{unique_id}_{original_name}"
                            fpath = os.path.join(UPLOAD_DIR, fname)
                            file.save(fpath)

                            url = request.host_url.rstrip('/') + f"/download/{fname}"
                            para.add_run("\n")
                            add_hyperlink(para, url, fname)

                            para = upload_doc.add_paragraph()
                            run = para.add_run(f"Row {row_index+1}, Col {col_index+1}: {fname}")
                            run.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)

                            ext = os.path.splitext(fname)[1].lower()
                            if ext in ['.png', '.jpg', '.jpeg']:
                                try:
                                    upload_doc.add_picture(fpath, width=Inches(4))
                                except:
                                    upload_doc.add_paragraph("(Error embedding image)")
                            elif ext == '.pdf':
                                pdf_to_images_with_fitz(fpath, upload_doc)
                                pdf_doc = pdf_to_docx_in_memory(fpath)
                                for p in pdf_doc.paragraphs:
                                    upload_doc.add_paragraph(p.text)
                            elif ext == '.docx':
                                upload_doc.add_paragraph("(Embedded DOCX content omitted for now)")
                            elif ext == '.txt':
                                with open(fpath, 'r', encoding='utf-8') as f:
                                    upload_doc.add_paragraph(f.read())
                            else:
                                upload_doc.add_paragraph(f"(Unsupported file type: {fname})")

                if not any_file_uploaded:
                    upload_doc.add_paragraph("No file uploaded for this question.")

            elif response_type == 'skip':
                response_doc.add_paragraph("Skipped by user.")
                upload_doc.add_paragraph("No file uploaded because question was skipped.")

            else:
                response_doc.add_paragraph("No valid response.")
                upload_doc.add_paragraph("No file uploaded due to no selected response type.")

        response_doc.save(response_path)
        upload_doc.save(uploads_path)
        
        if current_user.is_authenticated:
            new_submission = Submission(user_id=current_user.id, filename=filename)
            db.session.add(new_submission)
            db.session.commit()

        return jsonify({
            "success": True,
            "response_doc_url": f"/download/{os.path.basename(response_path)}",
            "upload_doc_url": f"/download/{os.path.basename(uploads_path)}"
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "message": f"Server error: {str(e)}"}), 500
    
   

    

# --------------------------------------
# Download Route
# --------------------------------------
@app.route('/download/<filename>')
def download(filename):
    path = os.path.join(UPLOAD_DIR, filename)
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    return "File not found.", 404






#---------------EXPORTPDF----------------
@app.route('/export-pdf')
def export_pdf():
    from asposewordscloud.models import PdfSaveOptionsData
    from asposewordscloud.models.requests import UploadFileRequest, SaveAsRequest, DownloadFileRequest

    filename = request.args.get('filename')
    if not filename:
        return "Missing 'filename' parameter.", 400

    response_docx = os.path.join(UPLOAD_DIR, f"{filename}_response.docx")
    uploads_docx = os.path.join(UPLOAD_DIR, f"{filename}_uploads.docx")

    if not os.path.exists(response_docx) or not os.path.exists(uploads_docx):
        return "One or both DOCX files not found.", 404

    # Auth with Aspose Cloud
    words_api = WordsApi(
        client_id=Config.ASPOSE_CLIENT_ID,
        client_secret=Config.ASPOSE_CLIENT_SECRET
    )

    response_pdf = os.path.join(UPLOAD_DIR, f"{filename}_response.pdf")
    uploads_pdf = os.path.join(UPLOAD_DIR, f"{filename}_uploads.pdf")

    try:
        for docx_path in [response_docx, uploads_docx]:
            docx_name = os.path.basename(docx_path)
            pdf_name = docx_name.replace(".docx", ".pdf")
            pdf_path = os.path.join(UPLOAD_DIR, pdf_name)

            # Upload to Aspose
            with open(docx_path, 'rb') as f:
                upload_req = UploadFileRequest(f, docx_name)
                words_api.upload_file(upload_req)

            # Convert to PDF
            save_options = PdfSaveOptionsData(file_name=pdf_name)
            save_req = SaveAsRequest(name=docx_name, save_options_data=save_options)
            words_api.save_as(save_req)

            # Download PDF
            download_req = DownloadFileRequest(pdf_name)
            result = words_api.download_file(download_req)

            with open(pdf_path, 'wb') as f:
                f.write(result)

        # Send back both PDF URLs
        return jsonify({
            "success": True,
            "response_pdf_url": f"/download/{filename}_response.pdf",
            "uploads_pdf_url": f"/download/{filename}_uploads.pdf"
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"PDF generation failed: {str(e)}", 500


  # Make sure this path is correct and accessible

@app.route('/generated-files/<username>')
def generated_reports(username):
    username = username.strip().replace(" ", "_").lower()
    user_dir = UPLOAD_DIR
    report_files = []

    for file in os.listdir(user_dir):
        if file.lower().startswith(username.lower() + "_") and file.endswith(".docx"):
            report_files.append(file)

    grouped = {}
    for file in report_files:
        if '_response.docx' in file:
            base = file.replace('_response.docx', '')
            grouped.setdefault(base, {})['response'] = file
        elif '_upload.docx' in file or '_uploads.docx' in file:
            base = file.replace('_upload.docx', '').replace('_uploads.docx', '')
            grouped.setdefault(base, {})['upload'] = file

    report_sets = list(grouped.values())
    return render_template("generatedreports.html", username=username.capitalize(), reports=report_sets)


# ------------------ Main Entry ------------------

if __name__ == '__main__':
    app.run(debug=True)
